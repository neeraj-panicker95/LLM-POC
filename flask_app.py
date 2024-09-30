import os
from flask import Flask, render_template, request, jsonify
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain_community.llms import OpenAI
from langchain.chains import RetrievalQA
from langchain.schema import Document
from docx import Document as DocxDocument

app = Flask(__name__)

# Initialize OpenAI LLM and Embeddings
openai_api_key = "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"  # Replace with your OpenAI API key
llm = OpenAI(openai_api_key=openai_api_key)
embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)

# A history list to maintain the last 5 conversation exchanges
conversation_histories = {}
conversation_id = "12345"

def load_docx_and_create_embeddings(docx_path):
    # Load the .docx file
    document = DocxDocument(docx_path)
    text_chunks = []

    # Extract text from each paragraph in the .docx file
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()  # Strip any leading/trailing whitespace
        if text:  # Ensure we only append non-empty paragraphs
            text_chunks.append(text)

    # If no text chunks were created, raise an error
    if not text_chunks:
        raise ValueError("No valid text chunks found. Ensure the .docx contains readable text.")

    # Convert text chunks to Document objects
    documents = [Document(page_content=text) for text in text_chunks]

    # Create embeddings for each text chunk (each chunk is a string)
    embeddings_list = embeddings.embed_documents(text_chunks)

    # If no embeddings were generated, raise an error
    if not embeddings_list:
        raise ValueError("No embeddings were generated. Check your API key and embedding function.")

    return documents, embeddings_list

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/chat', methods=['POST'])
def chat():
    user_message = request.json['message']
    conversation_id = request.json['conversation_id']

    # Initialize conversation history if it doesn't exist
    if conversation_id not in conversation_histories:
        conversation_histories[conversation_id] = []

    # Append the user's message to the conversation history
    conversation_histories[conversation_id].append(f"You: {user_message}")

    # Keep only the last 5 messages in history
    if len(conversation_histories[conversation_id]) > 5:
        conversation_histories[conversation_id].pop(0)  # Remove the oldest message

    # Create a custom prompt based on the user's message and the conversation history
    history_context = "\n".join(conversation_histories[conversation_id])  # Combine history into a single string
    custom_prompt = f"{history_context}\nUsing the context from the document, please answer the following question: {user_message}"

    # Query the vectorstore (from the .docx document) with the custom prompt
    response = qa_chain.run(custom_prompt)

    # Append the model's response to the conversation history
    conversation_histories[conversation_id].append(f"Bot: {response}")

    return jsonify({'response': response})

if __name__ == '__main__':
    # Load and embed your .docx file (replace with the path to your .docx file)
    docx_path = 'Y:\POC Works\myown-llm code\kungfu-panda.docx'  # Update this to your .docx's path
    text_chunks, embeddings_list = load_docx_and_create_embeddings(docx_path)
 
    # Create a FAISS vector store with the document embeddings
    vectorstore = FAISS.from_documents(text_chunks, embeddings)

    # Use the vectorstore to create a QA chain
    qa_chain = RetrievalQA.from_chain_type(llm, retriever=vectorstore.as_retriever())

    app.run(debug=True)
